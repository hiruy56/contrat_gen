import os
from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from fastapi.responses import StreamingResponse
from io import BytesIO

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["POST"],
    allow_headers=["*"],
)

def fill_contract_template(contract_details):
    template_path = os.path.abspath('tomp.docx')
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in contract_details.items():
            if f"{{{key}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in contract_details.items():
                        if f"{{{key}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

@app.post("/fill-contract")
async def fill_contract(contract_details: dict):
    filled_contract_stream = fill_contract_template(contract_details)
    return StreamingResponse(content=filled_contract_stream, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
